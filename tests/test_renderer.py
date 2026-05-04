from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.renderer import render_docx
from src.schemas import FinalReportContent, FinalReportItem


def _run_color(run) -> str:
    if run.font.color and run.font.color.rgb:
        return str(run.font.color.rgb)
    return ""


def test_render_docx_uses_business_heading_and_explicit_highlights(tmp_path: Path) -> None:
    report = FinalReportContent(
        report_title="监管政策周报",
        report_date="2026.04.26",
        items=[
            FinalReportItem(
                title="中国人民银行等八部门《金融产品网络营销管理办法》",
                url="",
                paragraph1="2026年4月24日，中国人民银行等八部门联合发布《金融产品网络营销管理办法》，自2026年9月30日起施行。",
                paragraph2="《办法》适用于平安集团及旗下各金融子公司：平安银行需清理数字营销矩阵；壹钱包需调整收银台页面。【已做专项解读】",
                paragraph2_mode="direct",
                importance="major",
                paragraph1_red_bold=["2026年9月30日"],
                paragraph2_black_bold=["平安银行", "壹钱包", "【已做专项解读】"],
            )
        ],
    )

    output = tmp_path / "weekly.docx"
    render_docx(report, output)

    doc = Document(str(output))
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]

    assert paragraphs[0].text == "监管政策周报"
    assert "合并版" not in paragraphs[0].text
    assert paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert paragraphs[1].text == "#2026.04.26"
    assert paragraphs[1].style.name == "Heading 2"
    assert paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert paragraphs[2].text == "1. 中国人民银行等八部门《金融产品网络营销管理办法》"
    assert paragraphs[2].style.name == "Heading 3"
    assert paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.LEFT

    runs = [run for paragraph in paragraphs for run in paragraph.runs if run.text]
    assert any(run.text == "2026年9月30日" and run.bold and _run_color(run) == "FF0000" for run in runs)
    assert any(run.text == "平安银行" and run.bold and _run_color(run) == "000000" for run in runs)
    assert any(run.text == "【已做专项解读】" and run.bold and _run_color(run) == "000000" for run in runs)


def test_render_docx_falls_back_to_red_bold_when_no_explicit_highlights(tmp_path: Path) -> None:
    report = FinalReportContent(
        report_title="监管政策周报",
        report_date="#2026.05.03",
        items=[
            FinalReportItem(
                title="测试机构《测试办法》",
                url="",
                paragraph1="测试机构发布《测试办法》。",
                paragraph2="对平安银行有直接影响。应开展存量业务排查。",
                paragraph2_mode="direct",
                importance="major",
            )
        ],
    )

    output = tmp_path / "fallback.docx"
    render_docx(report, output)

    doc = Document(str(output))
    runs = [run for paragraph in doc.paragraphs for run in paragraph.runs if run.text]

    assert any(run.bold and _run_color(run) == "FF0000" and "平安银行" in run.text for run in runs)
