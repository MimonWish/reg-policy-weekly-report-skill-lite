from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import RGBColor

from src.evolution import _paths_from_args, check_overfit, compare_reports, extract_docx_report, write_evolution_outputs


RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)


def _add_bold_run(paragraph, text: str, color=BLACK) -> None:
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = color


def _make_report(path: Path, p1: str, p2: str, *, extra_black: str = "") -> None:
    doc = Document()
    doc.add_heading("监管政策周报", level=1)
    doc.add_heading("#2026.05.10", level=2)
    doc.add_heading("1. 测试机构《测试政策》", level=3)

    p1_para = doc.add_paragraph()
    p1_para.add_run(p1)
    _add_bold_run(p1_para, "2026年6月1日", RED)

    p2_para = doc.add_paragraph()
    if extra_black:
        _add_bold_run(p2_para, extra_black, BLACK)
        p2_para.add_run("需同步落实。")
    else:
        p2_para.add_run(p2)
    doc.save(str(path))


def test_extract_and_compare_reports_emit_learning_proposals(tmp_path: Path) -> None:
    system_path = tmp_path / "system.docx"
    expert_path = tmp_path / "expert.docx"
    _make_report(
        system_path,
        "测试机构发布《测试政策》，自",
        "对平安银行有直接影响。应开展排查。",
        extra_black="平安银行",
    )
    _make_report(
        expert_path,
        "测试机构发布《测试政策》，明确新的监管要求，自",
        "平安银行、集团科技条线需同步落实。",
        extra_black="平安银行、集团科技条线",
    )

    extracted = extract_docx_report(system_path)
    assert extracted["report_title"] == "监管政策周报"
    assert extracted["report_date"] == "2026.05.10"
    assert extracted["item_count"] == 1
    assert extracted["style_summary"]["red_bold_count"] == 1
    assert extracted["style_summary"]["black_bold_count"] == 1

    result = compare_reports(system_path, expert_path)
    assert result["aggregate"]["matched_count"] == 1
    assert result["aggregate"]["entity_changed_count"] == 1
    proposal_types = {p["learning_type"] for p in result["learning_proposals"]}
    assert "impact_mapping" in proposal_types

    out = tmp_path / "evolution"
    registry = tmp_path / "evolution_rules.json"
    write_evolution_outputs(result, out, registry_path=registry)
    assert (out / "gap_analysis.json").exists()
    assert (out / "candidate_rules.json").exists()
    assert (out / "gap_analysis.md").read_text(encoding="utf-8").startswith("# 周报落差分析")
    assert (out / "learning_proposals.md").exists()
    assert (out / "evolution_patch_plan.md").exists()
    registry_text = registry.read_text(encoding="utf-8")
    assert "auto_high_confidence" in registry_text
    assert "集团科技条线" in registry_text


def test_check_overfit_detects_reused_body_paragraph(tmp_path: Path) -> None:
    candidate_path = tmp_path / "candidate.docx"
    historical_path = tmp_path / "historical.docx"
    reused = "这是一段非常具体的平安业务影响分析，用于模拟专家稿正文被系统稿复用的情况。"
    _make_report(candidate_path, "测试机构发布《测试政策》，自", reused)
    _make_report(historical_path, "历史机构发布《历史政策》，自", reused)

    report = check_overfit(candidate_path, [historical_path], threshold=0.72)
    assert report["hit_count"] >= 1
    assert report["max_similarity"] >= 0.72

    self_only = check_overfit(candidate_path, [candidate_path], threshold=0.72)
    assert self_only["hit_count"] == 0

    duplicate_path = tmp_path / "duplicate.docx"
    duplicate_path.write_bytes(candidate_path.read_bytes())
    duplicate_only = check_overfit(candidate_path, [duplicate_path], threshold=0.72)
    assert duplicate_only["hit_count"] == 0


def test_paths_from_args_skips_word_temp_files(tmp_path: Path) -> None:
    real = tmp_path / "real.docx"
    temp = tmp_path / "~$real.docx"
    real.write_bytes(b"placeholder")
    temp.write_bytes(b"placeholder")

    assert _paths_from_args((str(tmp_path),)) == [real]
