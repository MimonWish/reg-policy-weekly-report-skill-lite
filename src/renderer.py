from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from .schemas import FinalReportContent, FinalReportItem

RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)


def _add_run(paragraph, text: str, *, bold: bool = False, color: RGBColor | None = None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    if color is not None:
        run.font.color.rgb = color


def _detect_highlight_phrases(text: str) -> list[tuple[int, int]]:
    spans: list[tuple[int, int]] = []
    fixed_phrases = [
        "征求意见稿",
        "正式实施",
        "修订草案",
        "专项行动",
        "平安集团不直接适用",
    ]
    for phrase in fixed_phrases:
        for m in re.finditer(re.escape(phrase), text):
            spans.append((m.start(), m.end()))

    patterns = [
        r"对平安[^。；，]{2,40}有直接影响",
        r"后续将改变[^。；]{2,40}",
        r"(?:应立即|应尽快|应优先|应开展|应修订|应调整|应切换|可关注|可评估)[^。；]{0,36}",
    ]
    for pat in patterns:
        for m in re.finditer(pat, text):
            spans.append((m.start(), m.end()))

    spans.sort()
    merged: list[tuple[int, int]] = []
    for start, end in spans:
        if not merged or start > merged[-1][1]:
            merged.append((start, end))
        else:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
    return merged[:4]


def _find_phrase_spans(text: str, phrases: list[str]) -> list[tuple[int, int]]:
    spans: list[tuple[int, int]] = []
    for phrase in phrases:
        phrase = phrase.strip()
        if not phrase:
            continue
        start = 0
        while True:
            idx = text.find(phrase, start)
            if idx == -1:
                break
            spans.append((idx, idx + len(phrase)))
            start = idx + len(phrase)
    return spans


def _build_style_runs(
    text: str,
    *,
    red_bold_phrases: list[str] | None = None,
    black_bold_phrases: list[str] | None = None,
) -> list[tuple[str, str]]:
    """Return text chunks with style: normal, red_bold, or black_bold."""

    if not text:
        return []

    explicit = bool(red_bold_phrases or black_bold_phrases)
    marks: list[str | None] = [None] * len(text)

    for start, end in _find_phrase_spans(text, black_bold_phrases or []):
        for idx in range(start, end):
            marks[idx] = "black_bold"

    red_spans = _find_phrase_spans(text, red_bold_phrases or [])
    if not explicit:
        red_spans = _detect_highlight_phrases(text)
    for start, end in red_spans:
        for idx in range(start, end):
            marks[idx] = "red_bold"

    chunks: list[tuple[str, str]] = []
    cursor = 0
    while cursor < len(text):
        style = marks[cursor] or "normal"
        end = cursor + 1
        while end < len(text) and (marks[end] or "normal") == style:
            end += 1
        chunks.append((text[cursor:end], style))
        cursor = end
    return chunks


def _append_paragraph_with_highlight(
    doc: Document,
    text: str,
    *,
    red_bold_phrases: list[str] | None = None,
    black_bold_phrases: list[str] | None = None,
) -> None:
    if not text:
        return
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.35
    for chunk, style in _build_style_runs(
        text,
        red_bold_phrases=red_bold_phrases,
        black_bold_phrases=black_bold_phrases,
    ):
        if style == "red_bold":
            _add_run(paragraph, chunk, bold=True, color=RED)
        elif style == "black_bold":
            _add_run(paragraph, chunk, bold=True, color=BLACK)
        else:
            _add_run(paragraph, chunk, color=BLACK)


def _render_item(doc: Document, item: FinalReportItem, index: int) -> None:
    title_para = doc.add_heading(level=3)
    title_para.paragraph_format.space_before = Pt(10)
    _add_run(title_para, f"{index}. {item.title}", bold=True, color=BLACK)

    _append_paragraph_with_highlight(
        doc,
        item.paragraph1,
        red_bold_phrases=item.paragraph1_red_bold,
        black_bold_phrases=item.paragraph1_black_bold,
    )
    _append_paragraph_with_highlight(
        doc,
        item.paragraph2,
        red_bold_phrases=item.paragraph2_red_bold,
        black_bold_phrases=item.paragraph2_black_bold,
    )


def _format_report_date(report_date: str) -> str:
    return report_date if report_date.startswith("#") else f"#{report_date}"


def render_docx(report: FinalReportContent, output_path: Path) -> Path:
    doc = Document()
    title = doc.add_heading(report.report_title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = BLACK

    date_para = doc.add_heading(_format_report_date(report.report_date), level=2)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_after = Pt(12)
    for run in date_para.runs:
        run.font.color.rgb = BLACK

    for idx, item in enumerate(report.items, start=1):
        _render_item(doc, item, idx)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
